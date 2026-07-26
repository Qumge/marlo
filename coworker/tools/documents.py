"""Write .docx and .xlsx — the file formats a deliverable actually arrives in.

Marlo could produce markdown, csv and code, which is what a developer wants and
not what the people it is built for hand to anyone. "Write one document from a
folder of files" ended in a .md they then had to convert themselves, and that
one step is enough to decide an app was not meant for you.

Bundling python-docx and openpyxl is necessary but not sufficient: the agent's
only escape hatch is `shell`, the sidecar is a frozen PyInstaller binary with no
interpreter to import them from, and the user's own Mac has neither library. So
they arrive as tools.

The contract is markdown in, .docx out, because writing markdown is the one
authoring skill every model already has. Headings, paragraphs, bullets, numbered
lists and bold/italic runs map onto Word styles; anything richer is a document
someone should be editing in Word anyway.

Both tools write inside the workspace and refuse to escape it, the same boundary
the file tools enforce — with one difference that matters: these OVERWRITE. A
deliverable regenerated four times should not leave four files, and the approval
card names the path before any of this runs.
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Any

import aisuite as ai

_DOCX_SCHEMA = {
    "type": "function",
    "function": {
        "name": "write_docx",
        "description": (
            "Write a Word document (.docx) from markdown. Use this whenever the user asks "
            "for a document, report, letter, memo, or anything they will open in Word or "
            "send to someone — not a .md file, which they cannot use. Supports #/##/### "
            "headings, paragraphs, - bullets, 1. numbered lists, **bold** and *italic*."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Where to write it, relative to the workspace. Must end in .docx.",
                },
                "markdown": {
                    "type": "string",
                    "description": "The document body, as markdown.",
                },
                "title": {
                    "type": "string",
                    "description": "Optional. Sets the document's Title property (shown in Word's info panel).",
                },
            },
            "required": ["path", "markdown"],
        },
    },
}

_XLSX_SCHEMA = {
    "type": "function",
    "function": {
        "name": "write_xlsx",
        "description": (
            "Write an Excel workbook (.xlsx) from rows. Use this whenever the user asks for "
            "a spreadsheet, table, or list they will open in Excel or Numbers — not a .csv, "
            "which loses formatting and opens wrong on a Mac. The first row is treated as "
            "the header and is frozen and bolded."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Where to write it, relative to the workspace. Must end in .xlsx.",
                },
                "rows": {
                    "type": "array",
                    "description": "Rows of cells. The first row is the header.",
                    "items": {"type": "array", "items": {"type": "string"}},
                },
                "sheet_name": {
                    "type": "string",
                    "description": "Optional. The worksheet's tab name (default 'Sheet1').",
                },
            },
            "required": ["path", "rows"],
        },
    },
}

# **bold** and *italic*, non-greedy, on one line. Word runs are per-span, so the
# text is split on these and each piece carries its own formatting.
_RUN_SPLIT = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
_BULLET = re.compile(r"^\s*[-*+]\s+(.*)$")
_NUMBERED = re.compile(r"^\s*\d+[.)]\s+(.*)$")
_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")


def _add_runs(paragraph, text: str) -> None:
    for piece in _RUN_SPLIT.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            paragraph.add_run(piece[2:-2]).bold = True
        elif piece.startswith("*") and piece.endswith("*"):
            paragraph.add_run(piece[1:-1]).italic = True
        else:
            paragraph.add_run(piece)


def _resolve(root: Path, path: str, suffix: str) -> tuple[Path | None, str | None]:
    if not path.lower().endswith(suffix):
        return None, f"path must end in {suffix}"
    target = (root / path).resolve()
    try:
        target.relative_to(root)
    except ValueError:
        return None, "path escapes the workspace"
    return target, None


def document_tools(workspace: str) -> list:
    root = Path(workspace).resolve()

    def write_docx(path: str, markdown: str, title: str = "") -> dict[str, Any]:
        """Write a Word document from markdown."""
        target, err = _resolve(root, path, ".docx")
        if err:
            return {"error": err}

        try:
            from docx import Document
        except ImportError:  # pragma: no cover - only if packaging drops the dep
            return {"error": "python-docx is not available in this build"}

        doc = Document()
        if title:
            doc.core_properties.title = title

        for raw in (markdown or "").split("\n"):
            line = raw.rstrip()
            if not line.strip():
                continue
            if m := _HEADING.match(line):
                # Word's built-in Heading 1..6; level 0 is the document Title style,
                # which is not what "# " means in a body of markdown.
                doc.add_heading(m.group(2).strip(), level=min(len(m.group(1)), 6))
            elif m := _BULLET.match(line):
                _add_runs(doc.add_paragraph(style="List Bullet"), m.group(1))
            elif m := _NUMBERED.match(line):
                _add_runs(doc.add_paragraph(style="List Number"), m.group(1))
            else:
                _add_runs(doc.add_paragraph(), line)

        target.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(target))
        return {
            "path": str(target.relative_to(root)),
            "bytes": target.stat().st_size,
        }

    def write_xlsx(path: str, rows: list, sheet_name: str = "") -> dict[str, Any]:
        """Write an Excel workbook from rows; the first row is the header."""
        target, err = _resolve(root, path, ".xlsx")
        if err:
            return {"error": err}
        if not isinstance(rows, list) or not rows:
            return {"error": "rows must be a non-empty list of lists"}

        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font
            from openpyxl.utils import get_column_letter
        except ImportError:  # pragma: no cover
            return {"error": "openpyxl is not available in this build"}

        wb = Workbook()
        ws = wb.active
        if sheet_name:
            # Excel rejects these outright and silently truncates past 31 chars.
            ws.title = re.sub(r"[\\/*?:\[\]]", "-", sheet_name)[:31]

        widths: dict[int, int] = {}
        for r in rows:
            cells = r if isinstance(r, list) else [r]
            ws.append(["" if c is None else str(c) for c in cells])
            for i, c in enumerate(cells, start=1):
                widths[i] = max(widths.get(i, 0), len(str(c or "")))

        header = ws[1]
        for cell in header:
            cell.font = Font(bold=True)
        # Freeze it: a header that scrolls away is the difference between a
        # spreadsheet someone reads and one they close.
        ws.freeze_panes = "A2"
        for i, w in widths.items():
            ws.column_dimensions[get_column_letter(i)].width = min(max(w + 2, 8), 60)

        target.parent.mkdir(parents=True, exist_ok=True)
        wb.save(str(target))
        return {
            "path": str(target.relative_to(root)),
            "rows": len(rows),
            "bytes": target.stat().st_size,
        }

    docx_tool = ai.tool(
        write_docx,
        metadata=ai.ToolMetadata(
            category="files",
            # Writes a file, so it goes through the same approval gate as any edit.
            risk_level="medium",
            capabilities=["files", "documents"],
        ),
    )
    docx_tool.__coworker_schema__ = _DOCX_SCHEMA

    xlsx_tool = ai.tool(
        write_xlsx,
        metadata=ai.ToolMetadata(
            category="files",
            risk_level="medium",
            capabilities=["files", "documents"],
        ),
    )
    xlsx_tool.__coworker_schema__ = _XLSX_SCHEMA

    return [docx_tool, xlsx_tool]
