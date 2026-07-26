"""write_docx / write_xlsx — deliverables in the format they are handed over in.

Every assertion here reopens the file and reads what is inside it. Checking that
a path exists and has bytes would pass for a .docx containing one empty
paragraph, which is exactly the failure worth catching: the tool "worked", the
user opened it in Word, and there was nothing there.
"""
from __future__ import annotations

import pytest

from coworker.agents.base import AgentContext
from coworker.tools.documents import document_tools


@pytest.fixture
def tools(tmp_path):
    return {t.__name__: t for t in document_tools(str(tmp_path))}


@pytest.fixture
def docx(tools):
    return tools["write_docx"]


@pytest.fixture
def xlsx(tools):
    return tools["write_xlsx"]


# -- docx ---------------------------------------------------------------------

def test_markdown_becomes_word_structure(docx, tmp_path):
    res = docx(
        path="report.docx",
        markdown="# Quarterly review\n\nRevenue was **up**.\n\n- Asia grew\n- Europe flat\n\n1. Hire\n2. Ship",
        title="Q3",
    )
    assert "error" not in res

    from docx import Document

    doc = Document(str(tmp_path / "report.docx"))
    styles = [p.style.name for p in doc.paragraphs if p.text.strip()]
    texts = [p.text for p in doc.paragraphs if p.text.strip()]

    assert "Quarterly review" in texts
    assert any(s.startswith("Heading") for s in styles), "'# ' must become a Word heading"
    assert "List Bullet" in styles, "'- ' must become a bullet, not a line starting with a dash"
    assert "List Number" in styles
    assert doc.core_properties.title == "Q3"

    # The bold run has to be a run with bold set — markdown asterisks left in the
    # text would render as literal ** in Word.
    body = next(p for p in doc.paragraphs if p.text.startswith("Revenue"))
    assert "**" not in body.text
    assert any(r.bold and r.text == "up" for r in body.runs)


def test_it_refuses_to_write_outside_the_workspace(docx):
    assert "escapes the workspace" in docx(path="../escape.docx", markdown="x")["error"]


def test_it_refuses_a_path_that_is_not_a_docx(docx):
    # Writing docx bytes to report.md would produce a file nothing can open.
    assert ".docx" in docx(path="report.md", markdown="x")["error"]


def test_nested_paths_are_created(docx, tmp_path):
    assert "error" not in docx(path="out/2026/report.docx", markdown="# Hi")
    assert (tmp_path / "out/2026/report.docx").is_file()


def test_rewriting_replaces_rather_than_accumulates(docx, tmp_path):
    docx(path="r.docx", markdown="# First")
    docx(path="r.docx", markdown="# Second")

    from docx import Document

    texts = [p.text for p in Document(str(tmp_path / "r.docx")).paragraphs if p.text.strip()]
    assert texts == ["Second"], "a deliverable regenerated twice must not contain both drafts"


def test_markdown_tables_become_word_tables(docx, tmp_path):
    """The first real run produced a contract summary whose every table arrived as
    literal rows of "| 客户 | 蓝湖设计 |". Eleven unit tests were green, because
    they only covered what had been implemented — models reach for tables
    constantly in this kind of work and nothing had ever asked for one."""
    docx(
        path="t.docx",
        markdown=(
            "# Contracts\n\n"
            "| Client | Ends | Fee |\n"
            "|--------|------|----:|\n"
            "| Blue   | 2026-02-28 | 60,000 |\n"
            "| Star   | 2026-03-01 | 180,000 |\n\n"
            "Follow up soon."
        ),
    )

    from docx import Document

    doc = Document(str(tmp_path / "t.docx"))
    assert len(doc.tables) == 1, "a markdown table must become a Word table, not lines of text"

    table = doc.tables[0]
    assert len(table.rows) == 3  # header + two
    assert [c.text for c in table.rows[0].cells] == ["Client", "Ends", "Fee"]
    assert [c.text for c in table.rows[1].cells] == ["Blue", "2026-02-28", "60,000"]
    assert table.rows[0].cells[0].paragraphs[0].runs[0].bold, "header row reads as a header"

    # And nothing pipe-shaped survives in the prose.
    body = "\n".join(p.text for p in doc.paragraphs)
    assert "|" not in body
    assert "----" not in body
    assert "Follow up soon." in body, "text after the table still lands"


def test_a_lone_pipe_line_is_not_mistaken_for_a_table(docx, tmp_path):
    # Without the separator row underneath, it is prose that happens to contain
    # pipes — turning it into a one-row table would mangle ordinary writing.
    docx(path="p.docx", markdown="Use grep | sort to sort the output.")

    from docx import Document

    doc = Document(str(tmp_path / "p.docx"))
    assert len(doc.tables) == 0
    assert "grep | sort" in "\n".join(p.text for p in doc.paragraphs)


def test_horizontal_rules_do_not_print_as_hyphens(docx, tmp_path):
    docx(path="h.docx", markdown="# A\n\n---\n\n# B")

    from docx import Document

    texts = [p.text.strip() for p in Document(str(tmp_path / "h.docx")).paragraphs]
    assert "---" not in texts, "a thematic break must not land as three hyphens mid-page"


# -- xlsx ---------------------------------------------------------------------

def test_rows_become_a_readable_sheet(xlsx, tmp_path):
    res = xlsx(
        path="sales.xlsx",
        rows=[["Region", "Revenue"], ["Asia", "120"], ["Europe", "98"]],
        sheet_name="Q3 Sales",
    )
    assert "error" not in res
    assert res["rows"] == 3

    from openpyxl import load_workbook

    ws = load_workbook(str(tmp_path / "sales.xlsx")).active
    assert ws.title == "Q3 Sales"
    assert [c.value for c in ws[1]] == ["Region", "Revenue"]
    assert [c.value for c in ws[2]] == ["Asia", "120"]
    assert ws["A1"].font.bold, "the header row must read as a header"
    assert ws.freeze_panes == "A2", "a header that scrolls away is a spreadsheet nobody reads"


def test_columns_are_sized_to_their_contents(xlsx, tmp_path):
    # Not a fixed number: the first draft asserted `width > 8` and failed because
    # "Region" plus padding lands exactly on the floor. What matters is that a
    # long column ends up wider than a short one — a sheet whose text is all
    # #### is one nobody can read.
    xlsx(
        path="w.xlsx",
        rows=[["ID", "Description"], ["1", "A considerably longer piece of text here"]],
    )

    from openpyxl import load_workbook

    ws = load_workbook(str(tmp_path / "w.xlsx")).active
    assert ws.column_dimensions["B"].width > ws.column_dimensions["A"].width


def test_sheet_names_excel_would_reject_are_repaired(xlsx, tmp_path):
    # Excel refuses these characters outright and truncates past 31 chars; passing
    # them through produces a file that will not open.
    xlsx(path="s.xlsx", rows=[["a"]], sheet_name="Q3/Q4: [draft]" + "x" * 40)

    from openpyxl import load_workbook

    title = load_workbook(str(tmp_path / "s.xlsx")).active.title
    assert len(title) <= 31
    assert not set(title) & set("\\/*?:[]")


def test_empty_rows_are_rejected_rather_than_written(xlsx):
    assert "error" in xlsx(path="s.xlsx", rows=[])


def test_it_refuses_to_write_outside_the_workspace_xlsx(xlsx):
    assert "escapes the workspace" in xlsx(path="../escape.xlsx", rows=[["a"]])["error"]


# -- wiring -------------------------------------------------------------------

def test_the_cowork_agent_actually_gets_these_tools(tmp_path):
    # Registering a capability nobody composes is the same as not having one.
    from coworker.agents.cowork import cowork_tool_factory

    names = {getattr(t, "__name__", "") for t in cowork_tool_factory(AgentContext(workspace=str(tmp_path)))}
    assert {"write_docx", "write_xlsx"} <= names
